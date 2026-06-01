"""
File parser module for extracting text content from various file formats.
Supports: .txt, .md, .pdf, .docx, .ipynb, .py
"""
import json
import logging
import os
import re
from collections import Counter
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_txt(file_path: str) -> str:
    """Parse plain text file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def parse_markdown(file_path: str) -> str:
    """Parse Markdown file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def sanitize_control_chars(text: str) -> str:
    """Remove or replace ASCII control characters in text.

    Common PDF ligature mappings (font-specific, but widely observed):
    - U+001C (FS) → "fi" ligature
    - U+001B (ESC) → "ff" ligature
    - U+001D (GS) → "fl" ligature
    - U+001E (RS) → "ffi" ligature
    - U+001F (US) → "ffl" ligature

    Remaining control characters are replaced with a space so word
    boundaries are preserved and the frontend markdown parser doesn't
    choke on invalid XML/HTML characters.
    """
    # Map known PDF ligature control characters back to letters.
    # These mappings are heuristic — different PDFs may use different
    # encodings — but they are correct for the vast majority of cases.
    text = text.replace('\x1c', 'fi')   # FS → fi
    text = text.replace('\x1b', 'ff')   # ESC → ff
    text = text.replace('\x1d', 'fl')   # GS → fl
    text = text.replace('\x1e', 'ffi')  # RS → ffi
    text = text.replace('\x1f', 'ffl')  # US → ffl

    # Replace all remaining C0 control characters (except \t, \n, \r)
    # with a single space.  This keeps word boundaries intact so
    # "Retrieve" + U+0015 + "M" stays "Retrieve M" instead of becoming
    # "RetrieveM".
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1a]', ' ', text)

    # Replace C1 control characters (U+0080-U+009F) with a space.
    # These are invalid in HTML/XML and will break frontend parsing
    # when they survive into markdown content.
    text = re.sub(r'[\x80-\x9f]', ' ', text)

    return text


def _clean_pdf_text(text: str) -> str:
    """Post-process extracted PDF text to improve readability.

    Fixes common artifacts:
    - Sanitize control characters (ligatures, encoding artifacts)
    - Merge mid-paragraph line breaks (Chinese text)
    - Remove page numbers and running headers/footers
    - Normalize whitespace
    - Collapse excessive blank lines
    """
    if not text or not text.strip():
        return ""

    # Sanitize control characters FIRST — downstream processing and
    # the frontend markdown parser both break on invalid characters.
    text = sanitize_control_chars(text)

    # Remove common header/footer patterns: standalone page numbers,
    # repeated running headers (e.g. "Chapter 1  Introduction" on every page)
    lines = text.split('\n')
    cleaned: list[str] = []

    for line in lines:
        stripped = line.strip()
        # Skip standalone page numbers
        if re.match(r'^\d{1,4}$', stripped):
            continue
        # Skip lines that are just "第X页" or similar
        if re.match(r'^第[一二三四五六七八九十\d]+页$', stripped):
            continue
        cleaned.append(line)

    text = '\n'.join(cleaned)

    # Merge broken Chinese lines:
    # A line that doesn't end with a sentence-ending char or punctuation
    # and the next line starts with a Chinese char -> merge them
    text = re.sub(
        r'(?<=[^\n。！？；：，、"—…》\)\s])\n(?=[一-鿿㐀-䶿])',
        '',
        text
    )

    # Merge lines ending with hyphen (English word break across lines)
    text = re.sub(r'-\n(?=[a-zA-Z])', '', text)

    # Normalize whitespace: collapse 3+ newlines into double newline (paragraph break)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove excessive spaces within lines
    text = re.sub(r'[ \t]{2,}', ' ', text)

    return text.strip()


def parse_pdf(file_path: str) -> str:
    """Parse PDF file and extract text content using pymupdf.

    Falls back gracefully: if text extraction yields very little text,
    the caller can detect this and route to OCR.
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        raise ImportError(
            "pymupdf is required for PDF parsing. Install it with: pip install pymupdf"
        )

    text_pages: list[str] = []
    doc = fitz.open(file_path)

    try:
        for page in doc:
            # Use "text" mode for best plain-text extraction with layout awareness
            text = page.get_text("text")
            if text and text.strip():
                text_pages.append(text.strip())
    finally:
        doc.close()

    raw_text = '\n\n'.join(text_pages)
    return _clean_pdf_text(raw_text)


# ── PDF → Markdown extraction (font-metadata-based) ───────────────────

def _is_mono_font(font_name: str) -> bool:
    """Check if font name suggests a monospace/typewriter face."""
    mono_keywords = ("courier", "mono", "consolas", "code", "source code",
                     "typewriter", "menlo", "dejavu sans mono", "liberation mono",
                     "fira code", "jetbrains", "cascadia", "inconsolata")
    return any(kw in font_name.lower() for kw in mono_keywords)


def _is_bullet_start(line: str) -> bool:
    """Check if a line starts with a bullet or numbered-list marker."""
    stripped = line.strip()
    # Common bullet characters
    if stripped[:1] in "•▪▸▹◦○■□▪" and (len(stripped) < 2 or stripped[1] in (" ", "\t")):
        return True
    # Dash bullet: "- " or "– " at start
    if stripped.startswith(("- ", "– ", "— ")):
        return True
    # Asterisk bullet: "* " (but not "**" which is bold)
    if stripped.startswith("* ") and not stripped.startswith("**"):
        return True
    # Numbered list: "1.", "1)", "(1)", "1 ", "1.1", "i.", "a)"
    import re
    if re.match(r'^[\d]+[.)]\s', stripped):
        return True
    if re.match(r'^\([\d]+\)\s', stripped):
        return True
    if re.match(r'^[ivxlcdm]+[.)]\s', stripped, re.IGNORECASE):
        return True
    if re.match(r'^[a-z][.)]\s', stripped):
        return True
    return False


def _format_spans(spans: list[dict]) -> list[str]:
    """Convert a list of PyMuPDF span dicts to Markdown inline text.

    Each span dict must have 'text', 'flags', 'font' (optional), 'size'.
    Consecutive spans with identical bold/italic/monospace state are merged
    before wrapping, so ``**word1** **word2**`` becomes ``**word1 word2**``.
    """
    # Group consecutive spans that share the same formatting state
    groups: list[tuple[tuple[bool, bool, bool], str]] = []
    cur_state: tuple[bool, bool, bool] | None = None
    cur_text: str = ""

    for sp in spans:
        text = sp.get("text", "")
        if not text:
            continue

        flags = sp.get("flags", 0)
        font = sp.get("font", "")
        is_bold = bool(flags & 16)
        is_italic = bool(flags & 2)
        is_mono = bool(flags & 8) or _is_mono_font(font)
        state = (is_bold, is_italic, is_mono)

        if state == cur_state:
            cur_text += text
        else:
            if cur_text:
                groups.append((cur_state, cur_text))  # type: ignore[arg-type]
            cur_state = state
            cur_text = text

    if cur_text:
        groups.append((cur_state, cur_text))  # type: ignore[arg-type]

    # Wrap each group
    parts: list[str] = []
    for (is_bold, is_italic, is_mono), text in groups:
        text = text.rstrip()
        if not text:
            continue

        if is_bold and is_italic:
            text = f"***{text}***"
        elif is_bold:
            text = f"**{text}**"
        elif is_italic:
            text = f"*{text}*"

        if is_mono:
            text = f"`{text}`"

        parts.append(text)
    return parts


def parse_pdf_markdown(file_path: str) -> str:
    """Parse PDF and reconstruct Markdown formatting from font metadata.

    Uses PyMuPDF's ``get_text("dict")`` API to read font size, weight,
    italic flag, monospace flag, and bounding-box position for every span.
    These signals are mapped back to Markdown:

    - Font size (document-level relative scaling) → headings (h1–h4)
    - Flags bit 4 (bold) → ``**bold**``
    - Flags bit 1 (italic) → ``*italic*``
    - Flags bit 3 / font name → `` `code` `` or fenced code blocks
    - Bullet characters + indentation → unordered/ordered list items

    Returns clean Markdown text suitable for frontend rendering.
    """
    try:
        import fitz
    except ImportError:
        raise ImportError("pymupdf is required. Install: pip install pymupdf")

    doc = fitz.open(file_path)

    # ── Pass 1: collect all font sizes across the document ──────────
    all_sizes: list[float] = []
    for page in doc:
        page_dict = page.get_text("dict")
        for block in page_dict["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    all_sizes.append(round(span["size"], 1))

    if not all_sizes:
        doc.close()
        return ""

    # Document-level body-size detection (mode of all font sizes)
    size_counts = Counter(all_sizes)
    # Ignore very rare sizes (< 0.5% of spans) — they're usually artifacts
    min_count = max(len(all_sizes) * 0.005, 2)
    common_sizes = {s for s, c in size_counts.items() if c >= min_count}
    body_size = max(common_sizes, key=lambda s: size_counts[s]) if common_sizes else Counter(all_sizes).most_common(1)[0][0]

    # Heading sizes: sizes > body * 1.1, sorted descending → h1..h4
    heading_candidates = sorted(
        [s for s in common_sizes if s > body_size * 1.1],
        reverse=True
    )[:4]
    size_to_level: dict[float, int] = {}
    for i, size in enumerate(heading_candidates):
        size_to_level[size] = i + 1

    # ── Pass 2: generate Markdown ──────────────────────────────────
    all_pages_md: list[str] = []

    for page in doc:
        page_dict = page.get_text("dict")
        page_lines: list[str] = []
        prev_was_code_block = False

        for block in page_dict["blocks"]:
            if block.get("type") != 0:
                # Image block — signal a gap, but don't insert excessive blanks
                if page_lines and page_lines[-1] != "":
                    page_lines.append("")
                continue

            block_lines = block.get("lines", [])
            if not block_lines:
                continue

            # Check if this entire block is monospace → fenced code block
            block_all_mono = True
            for line in block_lines:
                for span in line["spans"]:
                    flags = span.get("flags", 0)
                    font = span.get("font", "")
                    if not (bool(flags & 8) or _is_mono_font(font)):
                        block_all_mono = False
                        break
                if not block_all_mono:
                    break

            if block_all_mono and len(block_lines) >= 1:
                code_parts: list[str] = []
                for line in block_lines:
                    text = "".join(sp["text"] for sp in line["spans"]).rstrip('\n')
                    code_parts.append(text)
                if prev_was_code_block:
                    # Merge into previous code block (separated by a page break
                    # but part of the same listing)
                    page_lines[-1] = page_lines[-1].rstrip('\n') + '\n' + '\n'.join(code_parts)
                else:
                    page_lines.append("```\n" + "\n".join(code_parts) + "\n```")
                page_lines.append("")
                prev_was_code_block = True
                continue

            prev_was_code_block = False

            # Separate blocks with a blank line
            if page_lines and page_lines[-1] != "":
                page_lines.append("")

            for line in block_lines:
                spans = line["spans"]
                if not spans:
                    continue

                max_size = max(round(sp["size"], 1) for sp in spans)
                line_text_raw = "".join(sp["text"] for sp in spans).strip()

                heading_level = size_to_level.get(max_size)

                # Fallback heuristic: short line (< 200 chars) with large font
                if heading_level is None and max_size > body_size * 1.2:
                    ratio = max_size / body_size
                    if ratio > 2.2:
                        heading_level = 1
                    elif ratio > 1.7:
                        heading_level = 2
                    elif ratio > 1.35:
                        heading_level = 3

                if heading_level and len(line_text_raw) < 200:
                    formatted = "".join(_format_spans(spans)).strip()
                    if formatted:
                        page_lines.append(f"{'#' * heading_level} {formatted}")
                else:
                    formatted = "".join(_format_spans(spans))
                    stripped = formatted.strip()
                    if not stripped:
                        page_lines.append("")
                    elif _is_bullet_start(stripped):
                        page_lines.append(stripped)
                    else:
                        page_lines.append(formatted)

        # Clean trailing blanks
        while page_lines and page_lines[-1] == "":
            page_lines.pop()
        while page_lines and page_lines[0] == "":
            page_lines.pop(0)

        if page_lines:
            all_pages_md.append("\n".join(page_lines))

    doc.close()
    return "\n\n".join(all_pages_md)


def parse_ipynb(file_path: str) -> str:
    """Parse Jupyter notebook and extract text from all cells."""
    with open(file_path, 'r', encoding='utf-8') as f:
        notebook = json.load(f)

    parts: list[str] = []
    for cell in notebook.get('cells', []):
        source = ''.join(cell.get('source', []))
        cell_type = cell.get('cell_type', 'code')

        if cell_type == 'markdown':
            parts.append(source)
        elif cell_type == 'code':
            parts.append(f'```python\n{source}\n```')

    return '\n\n'.join(parts)


def parse_docx(file_path: str) -> str:
    """Parse Word .docx file and extract text from all paragraphs."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install it with: pip install python-docx"
        )

    doc = Document(file_path)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text
        if text.strip():
            # Preserve heading styles with markdown-like prefix
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                paragraphs.append('#' * min(level, 6) + ' ' + text)
            else:
                paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = '\t'.join(cell.text for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)

    return '\n\n'.join(paragraphs)


def parse_file(file_path: str) -> str:
    """
    Parse file based on extension and return text content.

    Args:
        file_path: Path to the file

    Returns:
        Extracted text content

    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file does not exist
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_ext = Path(file_path).suffix.lower()

    parsers = {
        '.txt': parse_txt,
        '.md': parse_markdown,
        '.pdf': parse_pdf,
        '.docx': parse_docx,
        '.ipynb': parse_ipynb,
        '.py': parse_txt,
    }

    parser = parsers.get(file_ext)
    if not parser:
        raise ValueError(f"Unsupported file type: {file_ext}. Supported types: {', '.join(parsers.keys())}")

    return parser(file_path)


def extract_spans_from_pdf(file_path: str) -> list:
    """Extract spans with character positions and font metadata from a PDF.

    Thin wrapper around pdf_markdown.span_extractor.extract_spans.
    Returns list of Span objects with text, char_start, char_end, font info,
    bbox coordinates, and page_number.
    """
    from pdf_markdown.span_extractor import extract_spans
    return extract_spans(file_path)


def is_text_garbled(text: str, min_sample: int = 40) -> bool:
    """Detect whether extracted PDF text is likely garbled (mojibake).

    Returns True when the text appears to be garbage rather than real content.
    Common causes: non-standard CMap/font encoding in Chinese PDFs that pymupdf
    cannot decode correctly.

    IMPORTANT: Call this on RAW pymupdf text, before _clean_pdf_text sanitizes
    control characters. Once sanitized, the signals this function looks for
    (control chars, PUA chars) are already removed.

    Heuristics applied:
    - Count characters by Unicode block (CJK, ASCII, PUA, controls)
    - If >5% of chars are in Private Use Area → garbled
    - If C0 control chars (excluding \\t\\n\\r) >0.3% → garbled
    - If C1 control chars (U+0080-U+009F) >0.1% → garbled
    - If text has CJK chars but at very low ratio (<3%) among non-ASCII → garbled
    - Empty or very short text is NOT considered garbled (handled by needs_ocr)
    """
    if not text or not text.strip():
        return False

    stripped = text.strip()
    if len(stripped) < min_sample:
        return False

    total = 0
    cjk = 0       # CJK Unified Ideographs + Extensions
    ascii_printable = 0
    pua = 0       # Private Use Area
    c0_control = 0  # C0 control chars (U+0000-U+001F, excluding tab/newline/CR)
    c1_control = 0  # C1 control chars (U+0080-U+009F)
    other_non_ascii = 0

    for ch in stripped:
        cp = ord(ch)
        total += 1
        if 0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF or 0x20000 <= cp <= 0x2A6DF:
            cjk += 1
        elif 0x0020 <= cp <= 0x007E:
            ascii_printable += 1
        elif 0xE000 <= cp <= 0xF8FF:
            pua += 1
        elif cp < 0x0020 and cp not in (9, 10, 13):
            c0_control += 1
        elif 0x0080 <= cp <= 0x009F:
            c1_control += 1
        else:
            other_non_ascii += 1

    non_ascii = total - ascii_printable

    # C0 control characters (ligatures, encoding artifacts). Any real PDF
    # with >0.3% control chars has font encoding issues that will corrupt
    # math notation and break frontend markdown parsing.
    if c0_control > 0 and (c0_control / max(total, 1)) > 0.003:
        return True

    # C1 control characters are never legitimate in extracted PDF text,
    # but a small number (up to 1%) is tolerable — they're usually isolated
    # math symbols whose loss doesn't justify a multi-hour OCR run.
    if c1_control > 0 and (c1_control / max(total, 1)) > 0.01:
        return True

    # If >5% of characters are in Private Use Area, the text is garbled.
    # Legitimate PDFs rarely use PUA characters; >5% signals broken encoding.
    if pua > 0 and (pua / max(total, 1)) > 0.05:
        return True

    # If the text contains CJK characters but they represent <3% of non-ASCII
    # characters, the non-ASCII chars are likely encoding artifacts rather than
    # real content.
    if cjk > 0 and non_ascii > 0 and (cjk / non_ascii) < 0.03:
        return True

    # If non-ASCII dominates (>50%) but contains zero CJK and few ASCII chars,
    # the bytes were likely decoded as random Latin-1/Extended characters.
    if non_ascii > (total * 0.50) and cjk == 0 and ascii_printable < (total * 0.40):
        return True

    return False


def is_scanned_pdf(file_path: str) -> bool:
    """Check if a PDF appears to be scanned (image-based, needs OCR).

    Returns True if text extraction yields very little text relative to page count,
    suggesting the PDF contains mostly images rather than embedded text.
    """
    try:
        import fitz
    except ImportError:
        return False

    doc = fitz.open(file_path)
    try:
        page_count = len(doc)
        total_chars = 0
        for page in doc:
            total_chars += len(page.get_text("text").strip())

        # If average chars per page is very low, likely a scanned PDF
        avg_chars = total_chars / max(page_count, 1)
        return avg_chars < 50
    finally:
        doc.close()


def extract_pdf_images(file_path: str, output_dir: str) -> list[dict]:
    """Extract embedded images from a PDF and save them to output_dir.

    Returns a list of dicts with keys: index, page, width, height, filename.
    Images are saved as PNG files named page{N}_img{M}.png.
    """
    import fitz

    os.makedirs(output_dir, exist_ok=True)
    images: list[dict] = []
    doc = fitz.open(file_path)
    try:
        img_index = 0
        for page_num in range(len(doc)):
            page = doc[page_num]
            for img in page.get_images(full=True):
                xref = img[0]
                try:
                    base = doc.extract_image(xref)
                except Exception as e:
                    logger.warning("Failed to extract image xref=%s: %s", xref, e)
                    continue
                img_bytes = base.get("image")
                if not img_bytes:
                    continue
                ext = base.get("ext", "png")
                filename = f"page{page_num}_img{img_index}.{ext}"
                filepath = os.path.join(output_dir, filename)
                with open(filepath, "wb") as f:
                    f.write(img_bytes)
                images.append({
                    "index": img_index,
                    "page": page_num,
                    "width": base.get("width", 0),
                    "height": base.get("height", 0),
                    "filename": filename,
                })
                img_index += 1
        return images
    finally:
        doc.close()


def get_file_info(file_path: str) -> dict:
    """Get basic file information."""
    stat = os.stat(file_path)
    return {
        'name': os.path.basename(file_path),
        'size': stat.st_size,
        'extension': Path(file_path).suffix.lower(),
    }
